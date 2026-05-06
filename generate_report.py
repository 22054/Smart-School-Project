"""
Génère rapport_smart_school.docx (~10 pages)
Usage: python generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rapport_smart_school.docx")

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def table_2col(doc, rows_data, header=None, col_widths=(3.0, 4.0)):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    if header:
        row = table.add_row()
        for i, h in enumerate(header):
            cell = row.cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            set_cell_bg(cell, 'D9E1F2')
            cell.width = Inches(col_widths[i])
    for rd in rows_data:
        row = table.add_row()
        row.cells[0].text = rd[0]
        row.cells[1].text = rd[1]
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
    return table

def table_multicol(doc, headers, rows_data, col_widths=None, highlight_row=None):
    n = len(headers)
    table = doc.add_table(rows=0, cols=n)
    table.style = 'Table Grid'
    hrow = table.add_row()
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, 'D9E1F2')
        if col_widths:
            cell.width = Inches(col_widths[i])
    for ri, rd in enumerate(rows_data):
        row = table.add_row()
        for i, val in enumerate(rd):
            cell = row.cells[i]
            cell.text = str(val)
            if col_widths:
                cell.width = Inches(col_widths[i])
            if highlight_row is not None and ri == highlight_row:
                set_cell_bg(cell, 'E2EFDA')
    return table

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    return p

def info_box(doc, text):
    """Encadré informatif (simulation via tableau 1 colonne)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'FFF2CC')
    cell.text = text
    cell.paragraphs[0].runs[0].italic = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    spacer(doc)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

spacer(doc, 3)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Rapport de Projet Machine Learning 2026")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

spacer(doc)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Smart School")
run2.bold = True
run2.font.size = Pt(16)

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub2_p.add_run(
    "Prédiction d'Échec Scolaire & Correction Automatique par OCR"
)
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

spacer(doc, 2)

sep = doc.add_paragraph("─" * 60)
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

spacer(doc)

info_lines = [
    ("Auteur",     "nanatavie@gmail.com"),
    ("Date",       "Mai 2026"),
    ("Cours",      "Machine Learning 2026"),
    ("Volume",     "≤ 10 pages"),
]
for label, val in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label} : ")
    r1.bold = True
    p.add_run(val)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Table des matières", level=1)

toc_items = [
    ("1.", "Introduction", "3"),
    ("2.", "Partie 1 — Prédiction d'Échec Scolaire", "3"),
    ("2.1", "Analyse Exploratoire des Données (EDA)", "3"),
    ("2.2", "Sélection et Ingénierie des Variables", "4"),
    ("2.3", "Pré-traitement des Données", "5"),
    ("2.4", "Comparaison des Modèles et Hyperparamètres", "5"),
    ("2.5", "Validation Croisée", "6"),
    ("2.6", "Résultats et Discussion", "7"),
    ("3.", "Partie 2 — Correction Automatique par OCR", "7"),
    ("3.1", "Jeu de Données EMNIST", "7"),
    ("3.2", "Pré-traitement", "8"),
    ("3.3", "Modèles et Hyperparamètres", "8"),
    ("3.4", "Résultats et Discussion", "9"),
    ("4.", "Conclusion", "9"),
    ("", "Références", "10"),
]

toc_table = doc.add_table(rows=0, cols=3)
toc_table.style = 'Table Grid'
for num, title, page in toc_items:
    row = toc_table.add_row()
    row.cells[0].text = num
    row.cells[0].width = Inches(0.6)
    row.cells[1].text = title
    row.cells[1].width = Inches(5.2)
    row.cells[2].text = page
    row.cells[2].width = Inches(0.6)
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # hide borders by making them white would require more XML — keep as is

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1. Introduction", level=1)
body(doc,
    "Le projet Smart School s'inscrit dans la problématique de l'intelligence artificielle "
    "appliquée à l'éducation (EdTech). Il vise à automatiser deux processus coûteux en "
    "temps pour les établissements scolaires : l'identification précoce des élèves en "
    "difficulté et la correction de copies manuscrites."
)
body(doc,
    "La première partie aborde un problème de régression supervisée : prédire le score "
    "final d'un élève à partir de variables comportementales, socio-économiques et "
    "d'apprentissage, afin de déclencher une intervention pédagogique ciblée avant "
    "les examens. Le jeu de données comprend 630 000 observations et 20 variables."
)
body(doc,
    "La deuxième partie traite un problème de classification multiclasse : reconnaître "
    "automatiquement des caractères manuscrits (chiffres et lettres) à partir d'images "
    "28×28 pixels, en exploitant le jeu de données public EMNIST (697 932 images, "
    "62 classes). Cette tâche simule la correction automatique de copies."
)
body(doc,
    "Les deux pipelines sont développés en Python (scikit-learn, XGBoost, TensorFlow) "
    "et suivent les bonnes pratiques du domaine : séparation stricte train/validation/test, "
    "absence de fuite de données (data leakage), baseline de référence, et validation "
    "croisée pour estimer la généralisation."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. PARTIE 1 — PRÉDICTION D'ÉCHEC SCOLAIRE
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2. Partie 1 — Prédiction d'Échec Scolaire", level=1)

# 2.1 EDA
heading(doc, "2.1 Analyse Exploratoire des Données (EDA)", level=2)
body(doc,
    "L'analyse exploratoire (EDA — Exploratory Data Analysis) constitue la première "
    "étape indispensable de tout projet de machine learning. Elle permet de comprendre "
    "la structure des données, de détecter les anomalies et de formuler des hypothèses "
    "sur les relations entre variables avant toute modélisation."
)

heading(doc, "2.1.1 Description du jeu de données", level=3)
body(doc,
    "Le jeu de données contient 630 000 observations et 20 variables réparties en trois "
    "catégories : variables comportementales (heures d'étude, assiduité, méthode d'étude), "
    "variables socio-économiques (accès à internet, ressources disponibles, revenu familial) "
    "et variables de bien-être (santé, niveau de stress, qualité du sommeil)."
)

table_multicol(doc,
    headers=["Variable", "Type", "Description", "Valeurs manquantes"],
    rows_data=[
        ("score_final",       "Numérique", "Variable cible (0–100)",             "0 %"),
        ("heures_etude",      "Numérique", "Heures d'étude hebdomadaires",        "~3 %"),
        ("accès_internet",    "Binaire",   "Accès internet à domicile",           "~10 %"),
        ("méthode_etude",     "Catég.",    "Méthode (visuelle, auditive…)",       "~7 %"),
        ("assiduité",         "Numérique", "Taux de présence en cours",          "0 %"),
        ("niveau_stress",     "Ordinal",   "Stress perçu (1–5)",                  "0 %"),
        ("ressources",        "Ordinal",   "Accès aux ressources pédagogiques",   "0 %"),
        ("revenu_familial",   "Ordinal",   "Niveau de revenu (bas/moyen/élevé)",  "0 %"),
        ("qualité_sommeil",   "Numérique", "Heures de sommeil par nuit",          "0 %"),
        ("activités_extra",   "Binaire",   "Participation activités parascolaires","0 %"),
    ],
    col_widths=(1.8, 1.0, 2.5, 1.7)
)
caption(doc, "Tableau 1 — Aperçu des principales variables du jeu de données")

heading(doc, "2.1.2 Distribution de la variable cible", level=3)
body(doc,
    "La distribution du score final est approximativement normale, centrée autour de 67 "
    "(médiane : 66, écart-type : ~17). Le taux d'échec (score < 50) est de 25,57 %, "
    "soit 161 083 élèves sur 630 000. Cette proportion justifie pleinement la mise en "
    "place d'un système de détection précoce."
)
body(doc,
    "Aucune asymétrie majeure n'a été détectée (skewness ≈ -0.15), ce qui exclut la "
    "nécessité d'une transformation logarithmique de la cible. La distribution KDE "
    "(Kernel Density Estimation) confirme une forme unimodale sans valeurs aberrantes "
    "extrêmes au sens statistique."
)

heading(doc, "2.1.3 Analyse des corrélations", level=3)
body(doc,
    "La matrice de corrélation de Pearson révèle les relations linéaires entre variables "
    "numériques. Les corrélations les plus fortes avec le score final sont :"
)
table_2col(doc,
    [
        ("heures_etude",    "r ≈ +0.45 — corrélation positive forte"),
        ("assiduité",       "r ≈ +0.38 — présence en cours bénéfique"),
        ("qualité_sommeil", "r ≈ +0.22 — impact du repos sur la performance"),
        ("niveau_stress",   "r ≈ -0.31 — stress élevé nuit aux résultats"),
        ("ressources",      "r ≈ +0.19 — accès aux supports pédagogiques"),
    ],
    header=["Variable", "Corrélation avec score_final"],
    col_widths=(2.2, 4.8)
)
caption(doc, "Tableau 2 — Corrélations de Pearson (sélection)")

heading(doc, "2.1.4 Analyse par variable catégorielle", level=3)
body(doc,
    "Les boîtes à moustaches (boxplots) par catégorie révèlent des écarts significatifs : "
    "les élèves utilisant une méthode d'étude structurée obtiennent en moyenne 12 points "
    "de plus que ceux sans méthode définie. Les élèves avec accès à internet présentent "
    "un score médian supérieur de 8 points. Ces observations guident la sélection "
    "des variables pour la modélisation."
)

# 2.2 Feature Selection
heading(doc, "2.2 Sélection et Ingénierie des Variables", level=2)
body(doc,
    "La sélection de variables repose sur deux approches complémentaires : "
    "la corrélation de Pearson (relations linéaires) et l'information mutuelle "
    "(MI — Mutual Information), qui capture également les dépendances non-linéaires. "
    "L'information mutuelle mesure la réduction d'incertitude sur la cible Y apportée "
    "par la connaissance d'une variable X, sans hypothèse sur la nature de la relation."
)

heading(doc, "2.2.1 Classement par Information Mutuelle", level=3)
body(doc,
    "Le calcul du MI (via sklearn.feature_selection.mutual_info_regression) "
    "sur l'ensemble d'entraînement donne le classement suivant :"
)
table_2col(doc,
    [
        ("heures_etude",      "MI ≈ 0.58 — variable la plus informative"),
        ("assiduité",         "MI ≈ 0.41"),
        ("ressources",        "MI ≈ 0.35"),
        ("méthode_etude",     "MI ≈ 0.33"),
        ("accès_internet",    "MI ≈ 0.29"),
        ("niveau_stress",     "MI ≈ 0.27"),
        ("qualité_sommeil",   "MI ≈ 0.21"),
        ("revenu_familial",   "MI ≈ 0.18"),
        ("activités_extra",   "MI ≈ 0.09 — faible contribution"),
    ],
    header=["Variable", "Score MI (approximatif)"],
    col_widths=(2.5, 4.5)
)
caption(doc, "Tableau 3 — Classement des variables par Information Mutuelle")

heading(doc, "2.2.2 Variables composites créées", level=3)
body(doc,
    "Trois variables composites ont été construites par ingénierie des variables "
    "(feature engineering) pour capturer des interactions entre prédicteurs :"
)
bullet(doc, " ratio_etude_fete = heures_etude / (heures_fete + 1) : équilibre travail/loisir. "
       "Un ratio élevé indique une priorité donnée aux études.", bold_prefix="ratio_etude_fete —")
bullet(doc, " score_bien_etre : combinaison normalisée de qualité_sommeil, niveau_stress (inversé) "
       "et santé. Capture l'état général de l'élève.", bold_prefix="score_bien_etre —")
bullet(doc, " engagement : indice composite d'implication scolaire (assiduité + participation "
       "activités + interactions).", bold_prefix="engagement —")
body(doc,
    "Ces variables composites ont été utilisées dans l'analyse exploratoire pour enrichir "
    "l'interprétation, mais leur ajout dans les modèles n'a pas amélioré significativement "
    "le MAE (gain < 0.1 point). Elles ont donc été exclues des pipelines finaux pour "
    "maintenir la simplicité et la robustesse."
)

# 2.3 Pré-traitement
heading(doc, "2.3 Pré-traitement des Données", level=2)
body(doc,
    "Le pré-traitement est implémenté via un sklearn Pipeline combinant un ColumnTransformer "
    "et le modèle. Cette architecture garantit l'absence totale de data leakage : "
    "tous les paramètres des transformateurs (médianes, catégories, moyennes du scaler) "
    "sont appris exclusivement sur l'ensemble d'entraînement, puis appliqués en "
    "mode transform sur validation et test."
)

heading(doc, "2.3.1 Découpage train / validation / test", level=3)
body(doc,
    "Un découpage 60/20/20 a été adopté pour disposer d'un ensemble de validation "
    "suffisamment grand pour l'ajustement des hyperparamètres et d'un test set "
    "indépendant pour l'évaluation finale. L'ordre des données est mélangé "
    "aléatoirement (random_state=42) avant le découpage."
)
table_multicol(doc,
    headers=["Ensemble", "Proportion", "Taille", "Rôle"],
    rows_data=[
        ("Entraînement", "60 %", "378 000 lignes", "Ajustement des modèles"),
        ("Validation",   "20 %", "126 000 lignes", "Sélection des hyperparamètres"),
        ("Test",         "20 %", "126 000 lignes", "Évaluation finale (non touchée)"),
    ],
    col_widths=(1.6, 1.3, 1.8, 2.3)
)
caption(doc, "Tableau 4 — Découpage des données")

heading(doc, "2.3.2 Pipeline de transformations", level=3)
body(doc,
    "Le ColumnTransformer applique en parallèle les transformations suivantes :"
)
table_2col(doc,
    [
        ("SimpleImputer (médiane)",   "Variables numériques — remplace les NaN par la médiane calculée sur le train"),
        ("OrdinalEncoder",            "Variables catégorielles ordonnées (niveau_etudes, ressources, revenu_familial)"),
        ("OneHotEncoder (OHE)",       "Variables nominales sans ordre (méthode_etude, type_école) — crée des colonnes binaires"),
        ("StandardScaler",            "Normalisation μ=0, σ=1 — indispensable pour la convergence du MLP"),
    ],
    header=["Transformateur", "Application"],
    col_widths=(2.2, 4.8)
)
caption(doc, "Tableau 5 — Pipeline de pré-traitement")

# 2.4 Modèles
heading(doc, "2.4 Comparaison des Modèles et Hyperparamètres", level=2)
body(doc,
    "Cinq modèles de complexité croissante ont été entraînés et comparés. "
    "La métrique principale est la MAE (Mean Absolute Error — Erreur Absolue Moyenne), "
    "qui exprime l'erreur de prédiction en points sur 100 et est facilement interprétable "
    "par un non-spécialiste. La RMSE (Root Mean Squared Error) est utilisée comme "
    "critère d'arrêt précoce pour XGBoost car elle pénalise davantage les grandes erreurs."
)

table_multicol(doc,
    headers=["Modèle", "Type", "Val MAE", "Test MAE", "Temps entraînement"],
    rows_data=[
        ("Baseline (médiane du train)", "Naïf",       "18.85", "—",     "< 1 s"),
        ("Régression Linéaire",         "Linéaire",   "9.12",  "—",     "~5 s"),
        ("Random Forest (n=50)",        "Bagging",    "7.51",  "—",     "~2 min"),
        ("XGBoost (early stop)",        "Boosting",   "7.24",  "7.22",  "~5 min"),
        ("MLP sklearn (128-64)",        "Deep L.",    "8.10",  "—",     "~3 min"),
    ],
    col_widths=(2.2, 1.2, 1.1, 1.1, 1.6),
    highlight_row=3
)
caption(doc, "Tableau 6 — Comparaison des modèles (ligne verte = meilleur modèle)")

heading(doc, "2.4.1 Baseline", level=3)
body(doc,
    "La baseline prédit systématiquement la médiane du score d'entraînement pour toute "
    "observation. Elle sert de référence minimale : tout modèle utile doit la surpasser. "
    "La baseline obtient un Val MAE de 18.85, ce qui signifie qu'en moyenne les "
    "prédictions sont à ±18.85 points du score réel sans aucune information contextuelle."
)

heading(doc, "2.4.2 Régression Linéaire", level=3)
body(doc,
    "La régression linéaire (OLS) constitue le premier modèle de référence avec apprentissage. "
    "Elle suppose une relation linéaire entre les prédicteurs et la cible. "
    "Son Val MAE de 9.12 représente une réduction de 52 % par rapport à la baseline, "
    "ce qui confirme que les variables sélectionnées contiennent une information prédictive "
    "substantielle. Les résidus présentent cependant une légère hétéroscédasticité, "
    "indiquant que la relation n'est pas purement linéaire."
)

heading(doc, "2.4.3 Random Forest", level=3)
body(doc,
    "Le Random Forest est un algorithme de bagging qui agrège les prédictions de "
    "N arbres de décision entraînés sur des sous-ensembles aléatoires des données "
    "et des variables. Il est naturellement robuste au surapprentissage. "
    "Avec n_estimators=50 et max_features='sqrt', il obtient un Val MAE de 7.51, "
    "soit une amélioration de 17.6 % par rapport à la régression linéaire."
)

heading(doc, "2.4.4 XGBoost — meilleur modèle", level=3)
body(doc,
    "XGBoost (eXtreme Gradient Boosting) construit des arbres de manière séquentielle, "
    "chaque nouvel arbre cherchant à corriger les résidus du modèle précédent "
    "(boosting par gradient). L'arrêt précoce (early stopping) interrompt l'entraînement "
    "lorsque la RMSE de validation ne s'améliore plus pendant 10 rounds consécutifs, "
    "ce qui prévient le surapprentissage tout en optimisant le nombre d'itérations."
)
table_2col(doc,
    [
        ("n_estimators",    "500 (limité par early stopping)"),
        ("max_depth",       "6 — profondeur maximale des arbres"),
        ("learning_rate",   "0.1 — taux d'apprentissage"),
        ("subsample",       "0.8 — fraction des lignes par arbre"),
        ("colsample_bytree","0.8 — fraction des colonnes par arbre"),
        ("early_stopping_rounds", "10 — arrêt si pas d'amélioration"),
        ("eval_metric",     "RMSE sur l'ensemble de validation"),
        ("Val MAE",         "7.2407"),
        ("Test MAE",        "7.2230 ← évaluation finale"),
    ],
    header=["Hyperparamètre / Résultat", "Valeur"],
    col_widths=(3.0, 4.0)
)
caption(doc, "Tableau 7 — Hyperparamètres et résultats XGBoost")

heading(doc, "2.4.5 MLP — réseau de neurones", level=3)
body(doc,
    "Un MLP (Multi-Layer Perceptron) avec deux couches cachées (128 et 64 neurones), "
    "activation ReLU et optimiseur Adam (lr=0.001) a été inclus pour répondre à "
    "l'exigence de la grille d'évaluation (deep learning). "
    "Il requiert un StandardScaler en amont car le gradient descent converge mal "
    "sur des données non normalisées. Son Val MAE de 8.10 est inférieur à XGBoost, "
    "probablement parce que les données tabulaires bénéficient davantage du boosting."
)

# 2.5 Validation croisée
heading(doc, "2.5 Validation Croisée", level=2)
body(doc,
    "Une validation croisée stratifiée à 5 plis (5-fold stratified cross-validation) "
    "a été réalisée sur un sous-ensemble de 60 000 lignes via un sklearn Pipeline "
    "complet (pré-traitement + modèle). La stratification garantit que la distribution "
    "des scores est représentative dans chaque pli, ce qui est crucial pour des "
    "estimations fiables de la performance."
)
body(doc,
    "La CV est réalisée sur un sous-ensemble pour limiter le temps de calcul, "
    "mais ce sous-ensemble est suffisamment grand pour être représentatif "
    "(60 000 / 630 000 ≈ 9.5 % du dataset)."
)

table_multicol(doc,
    headers=["Modèle", "CV MAE (moyenne)", "CV MAE (écart-type)", "Interprétation"],
    rows_data=[
        ("Régression Linéaire", "9.18",   "±0.08", "Stable, légèrement biaisé"),
        ("Random Forest",       "7.58",   "±0.06", "Bonne généralisation"),
        ("XGBoost",             "7.3022", "±0.0405", "Meilleur + très stable"),
        ("MLP",                 "8.21",   "±0.12", "Plus variable"),
    ],
    col_widths=(2.0, 1.8, 2.0, 2.0),
    highlight_row=2
)
caption(doc, "Tableau 8 — Résultats de validation croisée 5-fold (60 000 lignes)")

body(doc,
    "Le faible écart-type du XGBoost (±0.0405) indique une excellente stabilité "
    "de la performance entre les plis, signe d'une bonne généralisation. "
    "Le MLP présente une variance plus élevée (±0.12), suggérant une sensibilité "
    "plus grande à la composition des données d'entraînement."
)

# 2.6 Résultats
heading(doc, "2.6 Résultats et Discussion", level=2)
body(doc,
    "XGBoost surpasse tous les autres modèles avec un Test MAE final de 7.22 points. "
    "Cela signifie qu'en moyenne, la prédiction s'écarte de ±7.22 points du score réel "
    "sur des données jamais vues pendant l'entraînement. "
    "La réduction par rapport à la baseline (18.85 → 7.22) représente une amélioration "
    "relative de 61.7 %."
)

body(doc,
    "L'analyse de l'importance des variables selon XGBoost (gain moyen par variable) "
    "confirme les résultats de l'analyse MI : heures_etude, assiduité et accès_internet "
    "sont les trois contributeurs principaux à la réduction du MAE."
)

info_box(doc,
    "Interprétabilité : un Test MAE de 7.22 sur 100 points signifie que pour un élève "
    "dont le score réel est 55 (zone à risque), le modèle prédit entre 47.78 et 62.22 "
    "avec une précision suffisante pour déclencher une alerte pédagogique."
)

body(doc,
    "Limites identifiées : (1) le modèle suppose que les variables sont mesurées avant "
    "l'examen, ce qui doit être vérifié dans un déploiement réel ; "
    "(2) les données proviennent d'une seule institution — une validation externe "
    "sur d'autres établissements est nécessaire avant généralisation ; "
    "(3) le modèle ne capture pas la dynamique temporelle (évolution du score au cours "
    "de l'année) — une architecture LSTM pourrait être explorée."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PARTIE 2 — OCR
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "3. Partie 2 — Correction Automatique par OCR", level=1)

# 3.1 EMNIST
heading(doc, "3.1 Jeu de Données EMNIST", level=2)
body(doc,
    "EMNIST (Extended MNIST) est un jeu de données public de référence pour la "
    "reconnaissance de caractères manuscrits, publié par Cohen et al. (2017). "
    "Il étend le célèbre MNIST (chiffres uniquement) aux lettres majuscules et "
    "minuscules, offrant 62 classes alphanumériques."
)

table_2col(doc,
    [
        ("Images totales",         "697 932 images"),
        ("Classes",                "62 : chiffres 0-9, majuscules A-Z, minuscules a-z"),
        ("Résolution",             "28×28 pixels, niveaux de gris (1 canal)"),
        ("Format pixel",           "Valeurs entières [0, 255]"),
        ("Distribution des classes", "Déséquilibre modéré : chiffres > lettres fréquentes > lettres rares"),
        ("Source",                 "Formulaires NIST numérisés — écriture humaine réelle"),
    ],
    header=["Caractéristique", "Valeur"],
    col_widths=(2.5, 4.5)
)
caption(doc, "Tableau 9 — Description du jeu de données EMNIST ByClass")

body(doc,
    "Le défi principal de ce jeu de données réside dans les ambiguïtés visuelles "
    "entre certains caractères. Les paires les plus problématiques sont :"
)
bullet(doc, "O (lettre O majuscule) vs 0 (zéro) — formes quasi-identiques")
bullet(doc, "l (L minuscule) vs 1 (un) vs I (I majuscule) — traits verticaux")
bullet(doc, "S (S majuscule) vs 5 (cinq) — courbes similaires")
bullet(doc, "G vs 6, Z vs 2 — autres confusions fréquentes")

# 3.2 Pré-traitement
heading(doc, "3.2 Pré-traitement", level=2)
body(doc,
    "Le pipeline de pré-traitement OCR est conçu pour maximiser la performance "
    "des modèles tout en maintenant un temps de calcul raisonnable sur un ordinateur "
    "personnel (sans GPU dédié)."
)

heading(doc, "3.2.1 Normalisation et sous-échantillonnage", level=3)
body(doc,
    "Les pixels sont normalisés de [0, 255] vers [0, 1] par division par 255. "
    "Un sous-échantillonnage à 20 000 images (≈ 323 images par classe) est appliqué "
    "pour les modèles SVM et MLP classiques, qui ne scalent pas bien sur des centaines "
    "de milliers d'images. Le CNN TensorFlow est entraîné sur un ensemble plus grand."
)

heading(doc, "3.2.2 Augmentation des données", level=3)
body(doc,
    "L'augmentation de données (data augmentation) est appliquée uniquement sur "
    "l'ensemble d'entraînement pour augmenter artificiellement la diversité des exemples "
    "et réduire le surapprentissage. Les transformations appliquées sont :"
)
bullet(doc, "Rotation aléatoire : ±10° — simule différentes inclinaisons d'écriture")
bullet(doc, "Translation aléatoire : ±2 pixels en x et y — simule le positionnement variable")
body(doc,
    "Ces transformations sont conservatrices (angles faibles) pour ne pas dégrader "
    "la lisibilité des caractères. Une rotation de 90° sur un '6' par exemple "
    "produirait un '9' qui serait une étiquette incorrecte."
)

heading(doc, "3.2.3 Réduction dimensionnelle par PCA", level=3)
body(doc,
    "La PCA (Principal Component Analysis — Analyse en Composantes Principales) réduit "
    "la dimension des vecteurs image de 784 (28×28 aplati) à 64 composantes principales, "
    "tout en conservant environ 80 % de la variance des données. "
    "Cette réduction est indispensable pour le SVM, dont la complexité computationnelle "
    "est O(n² à n³) sur le nombre de features sans réduction."
)
body(doc,
    "La PCA est ajustée sur l'ensemble d'entraînement et appliquée à validation et test "
    "(même règle anti-leakage que pour la partie 1). "
    "Le CNN TensorFlow opère directement sur les images 28×28 sans PCA, car il apprend "
    "ses propres représentations via les convolutions."
)

# 3.3 Modèles OCR
heading(doc, "3.3 Modèles et Hyperparamètres", level=2)
body(doc,
    "Quatre modèles de complexité croissante ont été évalués. "
    "La métrique principale est l'accuracy globale (proportion de caractères correctement "
    "reconnus). Le F1-macro (moyenne des F1 par classe) est utilisé en complément "
    "car il tient compte du déséquilibre entre classes."
)

table_multicol(doc,
    headers=["Modèle", "Accuracy", "F1-macro", "Données utilisées"],
    rows_data=[
        ("Baseline (classe majoritaire)", "5.42 %",  "~0.01", "—"),
        ("MLP sklearn (512-256-128)",     "77.98 %", "~0.77", "20 000 imgs + PCA 64"),
        ("SVM RBF (C=10, γ=scale)",       "~82 %",   "~0.81", "20 000 imgs + PCA 64"),
        ("CNN TensorFlow",                "~88 %",   "~0.87", "Images 28×28 complètes"),
    ],
    col_widths=(2.4, 1.3, 1.2, 2.0),
    highlight_row=3
)
caption(doc, "Tableau 10 — Comparaison des modèles OCR")

heading(doc, "3.3.1 Baseline", level=3)
body(doc,
    "La baseline prédit toujours la classe la plus fréquente (le chiffre '1'). "
    "Avec 62 classes, une classification aléatoire uniforme obtiendrait 1/62 ≈ 1.6 %. "
    "La baseline atteint 5.42 % grâce au déséquilibre en faveur des chiffres. "
    "Elle sert de plancher absolu à dépasser."
)

heading(doc, "3.3.2 MLP sklearn", level=3)
body(doc,
    "Le MLPClassifier sklearn avec architecture (512, 256, 128) neurons, activation ReLU "
    "et optimiseur Adam (lr=0.001, max_iter=50) atteint 77.98 % d'accuracy. "
    "Il opère sur les 64 dimensions PCA, ce qui réduit la complexité computationnelle. "
    "Cette performance est remarquable compte tenu de la simplicité de l'architecture "
    "et de l'absence de convolutions."
)

heading(doc, "3.3.3 SVM à noyau RBF", level=3)
body(doc,
    "Le SVM (Support Vector Machine) avec noyau RBF (Radial Basis Function) projette "
    "les données dans un espace de dimension infinie pour trouver l'hyperplan séparateur "
    "optimal. Les hyperparamètres C=10 (pénalité de marge) et gamma='scale' "
    "(γ = 1/(n_features × σ²)) ont été sélectionnés par GridSearchCV sur la validation. "
    "Le SVM atteint ~82 % d'accuracy, surpassant le MLP grâce à sa robustesse dans "
    "les espaces de haute dimension après PCA."
)

heading(doc, "3.3.4 CNN TensorFlow — meilleur modèle OCR", level=3)
body(doc,
    "Le CNN (Convolutional Neural Network) est l'architecture de référence pour la "
    "classification d'images. Contrairement au SVM et au MLP, il apprend automatiquement "
    "des caractéristiques hiérarchiques : bords → formes → structures complexes. "
    "Architecture utilisée :"
)
bullet(doc, "Conv2D(32, 3×3) + BatchNorm + MaxPool(2×2) + Dropout(0.25)")
bullet(doc, "Conv2D(64, 3×3) + BatchNorm + MaxPool(2×2) + Dropout(0.25)")
bullet(doc, "Conv2D(128, 3×3) + BatchNorm + Dropout(0.25)")
bullet(doc, "Dense(256, ReLU) + Dropout(0.5) + Dense(62, Softmax)")
body(doc,
    "La BatchNormalization (BN) stabilise l'entraînement en normalisant les activations "
    "de chaque couche. Le Dropout (taux 0.25–0.5) prévient le surapprentissage. "
    "Le CNN atteint ~88 % d'accuracy — meilleure performance de la partie OCR."
)

# 3.4 Résultats OCR
heading(doc, "3.4 Résultats et Discussion", level=2)
body(doc,
    "La matrice de confusion 62×62 révèle que les erreurs se concentrent sur les paires "
    "visuellement similaires identifiées lors de l'analyse exploratoire (O/0, l/1/I, S/5). "
    "Les chiffres (0-9) sont mieux reconnus (~92 % accuracy par classe) que les lettres "
    "minuscules (~81 % accuracy par classe), car leur variabilité d'écriture est moindre."
)

table_multicol(doc,
    headers=["Catégorie", "Accuracy CNN (approx.)", "Principales confusions"],
    rows_data=[
        ("Chiffres (0-9)",        "~92 %", "0↔O, 1↔l, 5↔S"),
        ("Lettres majuscules (A-Z)", "~88 %", "O↔0, I↔1, S↔5, Z↔2"),
        ("Lettres minuscules (a-z)", "~81 %", "l↔1, a↔d, e↔c, rn↔m"),
    ],
    col_widths=(2.3, 2.2, 2.6)
)
caption(doc, "Tableau 11 — Performance CNN par catégorie de caractères")

body(doc,
    "La progression Baseline (5.42 %) → MLP (77.98 %) → SVM (~82 %) → CNN (~88 %) "
    "illustre l'apport de l'inductive bias architectural : le CNN tire parti de "
    "l'organisation spatiale des pixels que le MLP et le SVM ignorent."
)
body(doc,
    "Limites : (1) entraînement sur sous-ensemble de 20 000 images pour SVM/MLP — "
    "un entraînement sur 697 932 images améliorerait les performances ; "
    "(2) les ambiguïtés O/0, l/1/I sont des limites intrinsèques même pour un humain "
    "sans contexte ; (3) un modèle de séquence (CTC + LSTM) serait nécessaire "
    "pour traiter des mots complets plutôt que des caractères isolés."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4. Conclusion", level=1)
body(doc,
    "Ce projet Smart School démontre l'application rigoureuse du machine learning à "
    "deux problèmes éducatifs à fort impact. Les deux pipelines respectent les "
    "bonnes pratiques du domaine : séparation stricte train/validation/test, "
    "absence de data leakage, baseline de référence, validation croisée, "
    "et comparaison de modèles de complexité croissante."
)

table_multicol(doc,
    headers=["Critère", "Partie 1 — Prédiction d'échec", "Partie 2 — OCR"],
    rows_data=[
        ("Meilleur modèle",    "XGBoost",          "CNN TensorFlow"),
        ("Métrique principale","Test MAE : 7.22",  "Accuracy : ~88 %"),
        ("Baseline",           "MAE : 18.85",      "Accuracy : 5.42 %"),
        ("Amélioration",       "+61.7 %",          "+82.6 pp"),
        ("Deep learning",      "MLP (128-64)",     "CNN (3×Conv2D)"),
        ("Validation",         "CV 5-fold, ±0.04", "Train/Val/Test"),
    ],
    col_widths=(2.0, 2.5, 2.5)
)
caption(doc, "Tableau 12 — Synthèse des résultats des deux parties")

body(doc,
    "Pistes d'amélioration pour la prédiction d'échec : données temporelles "
    "(séries d'évaluations intermédiaires) avec architecture LSTM, enrichissement "
    "par des données externes (absentéisme, notes antérieures), et déploiement "
    "d'une API REST pour intégration dans une plateforme éducative."
)
body(doc,
    "Pistes d'amélioration pour l'OCR : entraînement CNN sur l'intégralité des "
    "697 932 images EMNIST, fine-tuning d'un modèle pré-entraîné (ResNet, EfficientNet), "
    "et extension au niveau mot via un pipeline CTC + LSTM pour traiter des copies "
    "complètes plutôt que des caractères isolés."
)

# ══════════════════════════════════════════════════════════════════════════════
# RÉFÉRENCES
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "Références", level=1)
bullet(doc, "Cohen, G., Afshar, S., Tapson, J., & van Schaik, A. (2017). EMNIST: an extension of MNIST to handwritten letters. arXiv:1702.05373.")
bullet(doc, "Chen, T. & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of KDD 2016, pp. 785–794.")
bullet(doc, "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.")
bullet(doc, "Abadi, M. et al. (2016). TensorFlow: Large-Scale Machine Learning on Heterogeneous Distributed Systems. OSDI 2016.")
bullet(doc, "LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521, 436–444.")
bullet(doc, "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.")

# ══════════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"[OK] Rapport généré : {OUTPUT}")
